#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_massive_document(output_path, title, sections):
    """
    Menciptakan dokumen .docx besar secara otomatis.
    output_path: Lokasi file .docx yang akan dihasilkan.
    title: Judul utama dokumen.
    sections: List of dict {'heading': '...', 'content': '...'}
    """
    doc = Document()
    
    # Judul Utama
    main_title = doc.add_heading(title, 0)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for section in sections:
        doc.add_heading(section['heading'], level=1)
        p = doc.add_paragraph(section['content'])
        
        # Contoh kustomisasi font
        run = p.runs[0]
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Tambahkan page break jika perlu (opsional)
        # doc.add_page_break()

    print(f"[*] Dokumen berhasil dibuat: {output_path}")
    doc.save(output_path)

if __name__ == "__main__":
    # Contoh penggunaan sederhana
    output = "private/docx_source/PMN_Generated_Document.docx"
    
    # Pastikan folder tujuan ada
    os.makedirs(os.path.dirname(output), exist_ok=True)
    
    sample_sections = [
        {
            "heading": "Bab 1: Pendahuluan",
            "content": "Ini adalah konten otomatis untuk bab pertama. Dalam skala besar, Anda bisa melakukan loop pada ribuan data di sini."
        },
        {
            "heading": "Bab 2: Analisis Sistem",
            "content": "Konten bab kedua yang dihasilkan secara terprogram menggunakan python-docx."
        }
    ]
    
    create_massive_document(output, "PMN ULTIMATE DOCUMENT", sample_sections)
